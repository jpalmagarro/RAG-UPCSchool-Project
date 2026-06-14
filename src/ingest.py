import os
import glob
import re
import unicodedata
import argparse
import logging
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))

from google.cloud import storage

import chromadb
import io
import subprocess
from docx import Document
from transformers import AutoModel, AutoTokenizer
from googleapiclient.discovery import build
from google.oauth2 import service_account
from googleapiclient.http import MediaIoBaseDownload

from rag.bm25_index import build_bm25_index
from rag.config import (
    bm25_path,
    COLLECTION_NAME,
    COLLECTION_NAME_V2,
    EMBED_MODEL_NAME,
)
from rag.embeddings import E5Embedder

logging.basicConfig(level=logging.INFO, format='%(levelname)s: %(message)s')

from dotenv import load_dotenv
load_dotenv()

BUCKET_NAME = os.getenv("GCP_BUCKET_NAME")
CHROMA_PATH = os.getenv("CHROMA_PATH", "./data/chroma_db")
RAW_DOCS_PATH = os.getenv("RAW_DOCS_PATH", "./data/raw")
DRIVE_FOLDER_ID = os.getenv("DRIVE_FOLDER_ID")
CREDENTIALS_FILE = "credentials.json"

TABLE_ROW_SEP = " || "

# --- UTILIDADES DE GCP ---
def download_folder_from_gcs(bucket_name, source_folder, destination_folder):
    storage_client = storage.Client()
    bucket = storage_client.bucket(bucket_name)
    blobs = bucket.list_blobs(prefix=source_folder)

    count = 0
    for blob in blobs:
        if blob.name.endswith("/"):
            continue
        relative_path = os.path.relpath(blob.name, source_folder)
        local_path = os.path.join(destination_folder, relative_path)
        os.makedirs(os.path.dirname(local_path), exist_ok=True)
        blob.download_to_filename(local_path)
        count += 1
    logging.info(f"Descargados {count} archivos de gs://{bucket_name}/{source_folder} a {destination_folder}")

def upload_folder_to_gcs(bucket_name, source_folder, destination_folder):
    storage_client = storage.Client()
    bucket = storage_client.bucket(bucket_name)

    count = 0
    for root, _, files in os.walk(source_folder):
        for file in files:
            local_path = os.path.join(root, file)
            relative_path = os.path.relpath(local_path, source_folder)
            blob_path = os.path.join(destination_folder, relative_path).replace("\\", "/")

            blob = bucket.blob(blob_path)
            blob.upload_from_filename(local_path)
            count += 1
    logging.info(f"Subidos {count} archivos de {source_folder} a gs://{bucket_name}/{destination_folder}")

def download_folder_from_drive(folder_id, destination_folder):
    """Descarga todos los archivos .docx de una carpet de Google Drive usando la API."""
    import google.auth
    SCOPES = ['https://www.googleapis.com/auth/drive.readonly']

    try:
        if os.path.exists(CREDENTIALS_FILE):
            logging.info("Usando credentials.json para autenticación (Modo Local).")
            creds = service_account.Credentials.from_service_account_file(CREDENTIALS_FILE, scopes=SCOPES)
        else:
            logging.info("credentials.json no encontrado. Usando identidad de Google Cloud (Modo Cloud Run).")
            creds, _ = google.auth.default(scopes=SCOPES)

        service = build('drive', 'v3', credentials=creds)

        query = f"'{folder_id}' in parents and trashed = false"
        results = service.files().list(q=query, fields="files(id, name)").execute()
        items = results.get('files', [])

        if not items:
            logging.warning("No se encontraron archivos en la carpet de Google Drive.")
            return True

        os.makedirs(destination_folder, exist_ok=True)

        count = 0
        for item in items:
            if not (item['name'].lower().endswith('.docx') or item['name'].lower().endswith('.doc')):
                continue

            file_id = item['id']
            file_name = item['name']
            file_path = os.path.join(destination_folder, file_name)

            request = service.files().get_media(fileId=file_id)
            fh = io.FileIO(file_path, 'wb')
            downloader = MediaIoBaseDownload(fh, request)
            done = False
            while done is False:
                status, done = downloader.next_chunk()
            count += 1
            logging.info(f"Descargado de Drive: {file_name}")

        logging.info(f"✅ Descargados {count} archivos de Word de Google Drive.")
        return True
    except Exception as e:
        logging.error(f"❌ Error al conectar con Google Drive API: {e}")
        return False

# --- PROCESAMIENTO DE TEXTO ---
def normalize_text(text):
    return re.sub(r"\s+", " ", text).strip()

def clean_text(text):
    text = unicodedata.normalize("NFKC", text)
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
    text = re.sub(r"[\u00ad\u2010-\u2015\u2212\uff0d]", "-", text)
    text = re.sub(r"[\u2018\u2019\u201a\u201b]", "'", text)
    text = re.sub(r"[\u201c\u201d\u201e\u201f]", '"', text)
    text = re.sub(r"[\u2022\u2023\u25cf\u25e6\u2043\u2219\u00b7]", "-", text)
    text = re.sub(r"(?m)^\s*\d{1,4}\s*$", "", text)
    text = re.sub(r"([.\-_=])\1{2,}", r"\1", text)
    text = re.sub(r"\s+([.,;:!?\)])", r"\1", text)
    text = re.sub(r"(?mi)^\s*(fecha|lugar|hora de inicio|hora de finalizacion|asistentes)\s*$", "", text)
    text = re.sub(r"(?mi)^\s*(firmado|firma).*$", "", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()

def is_bad_chunk(chunk, min_words=40, min_alpha_ratio=0.65):
    words = len(chunk.split())
    if words < min_words:
        return True
    alpha_ratio = sum(c.isalpha() for c in chunk) / max(len(chunk), 1)
    return alpha_ratio < min_alpha_ratio

def chunk_text_recursive(text, chunk_size=500, overlap=100, min_chunk_size=50):
    if not text:
        return []
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=overlap,
        separators=["\n\n", "\n", ". ", " ", ""],
        length_function=len,
    )
    raw_chunks = splitter.split_text(text)
    return [c for c in raw_chunks if len(c) >= min_chunk_size and not is_bad_chunk(c)]

def chunk_text_table_hybrid(text, chunk_size=500, overlap=100, min_chunk_size=50):
    """Párrafos → recursive splitter; filas de tabla (||) → un chunk por fila."""
    if not text:
        return []

    paragraphs: list[str] = []
    table_rows: list[str] = []

    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        if TABLE_ROW_SEP in line:
            table_rows.append(line)
        else:
            paragraphs.append(line)

    chunks: list[str] = []
    if paragraphs:
        chunks.extend(chunk_text_recursive("\n\n".join(paragraphs), chunk_size, overlap, min_chunk_size))

    for row in table_rows:
        if len(row) >= min_chunk_size and not is_bad_chunk(row):
            chunks.append(row)

    return chunks

def chunk_document(text, strategy="recursive", **kwargs):
    if strategy == "table_hybrid":
        return chunk_text_table_hybrid(text, **kwargs)
    return chunk_text_recursive(text, **kwargs)

def extract_docx(filepath):
    try:
        doc = Document(filepath)
    except Exception as e:
        logging.error(f"Error leyendo {filepath}: {e}")
        return None

    texts = []
    for paragraph in doc.paragraphs:
        text = normalize_text(paragraph.text)
        if text:
            texts.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_texts = []
            seen = set()
            for cell in row.cells:
                text = normalize_text(cell.text)
                if not text or text in seen:
                    continue
                seen.add(text)
                row_texts.append(text)
            if row_texts:
                texts.append(TABLE_ROW_SEP.join(row_texts))

    return "\n".join(texts).strip() or None

def extract_doc_antiword(filepath):
    try:
        result = subprocess.run(['antiword', filepath], capture_output=True, text=True)
        if result.returncode == 0:
            return result.stdout.strip() or None
        logging.error(f"Error procesando .doc con antiword {filepath}: {result.stderr}")
        return None
    except Exception as e:
        logging.warning(f"No se pudo usar antiword para {filepath}. ¿Está instalado? Error: {e}")
        return None

def deduplicate_filepaths(filepaths):
    """Un acta puede existir como .doc y .docx en Drive; preferimos .docx."""
    by_id: dict[str, str] = {}

    def priority(fp: str) -> int:
        low = fp.lower()
        if low.endswith(".docx"):
            return 0
        if low.endswith(".doc"):
            return 1
        return 2

    for fp in filepaths:
        doc_id = os.path.splitext(os.path.basename(fp))[0]
        if doc_id not in by_id:
            by_id[doc_id] = fp
            continue
        if priority(fp) < priority(by_id[doc_id]):
            logging.warning("Duplicado: %s sustituye a %s", os.path.basename(fp), os.path.basename(by_id[doc_id]))
            by_id[doc_id] = fp
        else:
            logging.warning("Duplicado ignorado: %s (ya usamos %s)", os.path.basename(fp), os.path.basename(by_id[doc_id]))

    return list(by_id.values())


def process_all_documents(raw_dir, chunk_strategy="recursive"):
    documents = []
    filepaths_docx = glob.glob(os.path.join(raw_dir, "*.docx"))
    filepaths_doc = glob.glob(os.path.join(raw_dir, "*.doc"))
    filepaths = deduplicate_filepaths(filepaths_docx + filepaths_doc)

    if not filepaths:
        logging.warning(f"No se encontraron archivos de Word en {raw_dir}")
        return []

    for filepath in filepaths:
        filename = os.path.basename(filepath)
        doc_id = os.path.splitext(filename)[0]

        if filepath.lower().endswith('.docx'):
            raw = extract_docx(filepath)
        else:
            raw = extract_doc_antiword(filepath)

        if not raw:
            continue
        cleaned = clean_text(raw)
        chunks = chunk_document(cleaned, strategy=chunk_strategy)
        if chunks:
            documents.append({"doc_id": doc_id, "chunks": chunks})

    return documents

def index_to_chroma(documents, chroma_path, collection_name):
    embedder = E5Embedder.from_pretrained(EMBED_MODEL_NAME)

    all_texts, all_doc_ids, all_ids = [], [], []
    for doc in documents:
        for i, chunk in enumerate(doc["chunks"]):
            all_texts.append(chunk)
            all_doc_ids.append(doc["doc_id"])
            all_ids.append(f"{doc['doc_id']}__c{i:04d}")

    if not all_texts:
        logging.warning("No hay textos para indexar.")
        return

    if len(all_ids) != len(set(all_ids)):
        dupes = len(all_ids) - len(set(all_ids))
        raise ValueError(f"IDs de chunk duplicados ({dupes}); revisa duplicados .doc/.docx en Drive.")

    logging.info(f"Generando embeddings para {len(all_texts)} chunks...")
    all_embeddings = embedder.embed_passages(all_texts)

    logging.info(f"Guardando en ChromaDB ({chroma_path}), colección {collection_name}...")
    client = chromadb.PersistentClient(path=chroma_path)
    try:
        client.delete_collection(collection_name)
    except Exception:
        pass

    collection = client.create_collection(name=collection_name, metadata={"hnsw:space": "cosine"})

    metadatas = [{"doc_id": doc_id, "chunk_id": cid} for doc_id, cid in zip(all_doc_ids, all_ids)]

    for i in range(0, len(all_texts), 5000):
        j = min(i + 5000, len(all_texts))
        collection.add(
            ids=all_ids[i:j],
            documents=all_texts[i:j],
            embeddings=all_embeddings[i:j],
            metadatas=metadatas[i:j],
        )
    logging.info(f"✅ Indexación completada. Total chunks: {collection.count()}")

    bm25 = build_bm25_index(all_texts, all_ids)
    out_path = bm25_path(chroma_path, collection_name=collection_name)
    bm25.save(out_path)
    logging.info(f"✅ BM25 persistido en {out_path} ({len(all_ids)} docs)")


INDEX_TARGETS = {
    "recursive": (COLLECTION_NAME, "recursive"),
    "table_hybrid": (COLLECTION_NAME_V2, "table_hybrid"),
}


def run_indexing(raw_dir: str, chroma_path: str, chunk_strategy: str) -> None:
    """Indexa una o ambas colecciones en el mismo chroma_path (sin pisar la otra)."""
    if chunk_strategy == "both":
        targets = [INDEX_TARGETS["recursive"], INDEX_TARGETS["table_hybrid"]]
    else:
        targets = [INDEX_TARGETS[chunk_strategy]]

    for collection_name, strategy in targets:
        logging.info("=== Indexando %s (strategy=%s) ===", collection_name, strategy)
        documents = process_all_documents(raw_dir, chunk_strategy=strategy)
        if not documents:
            logging.warning("Sin documentos para %s, salto.", collection_name)
            continue
        index_to_chroma(documents, chroma_path, collection_name)

def main():
    parser = argparse.ArgumentParser(description="Pipeline de Ingesta MLOps para COSORA")
    parser.add_argument("--download-raw", action="store_true", help="Descargar actas .docx desde GCP Bucket")
    parser.add_argument("--download-drive", action="store_true", help="Descargar actas .docx desde Google Drive API")
    parser.add_argument("--upload-db", action="store_true", help="Subir ChromaDB a GCP Bucket al finalizar")
    parser.add_argument(
        "--chunk-strategy",
        choices=["recursive", "table_hybrid", "both"],
        default="both",
        help="recursive→cosora_actas_e5 | table_hybrid→v2 | both→las dos colecciones",
    )
    parser.add_argument(
        "--collection",
        default=None,
        help="Forzar nombre de colección (solo con recursive o table_hybrid, no con both)",
    )
    args = parser.parse_args()

    if args.chunk_strategy == "both" and args.collection:
        logging.error("--collection no es compatible con --chunk-strategy both")
        exit(1)

    if args.collection:
        collection_name = args.collection
    elif args.chunk_strategy == "table_hybrid":
        collection_name = COLLECTION_NAME_V2
    elif args.chunk_strategy == "recursive":
        collection_name = os.getenv("CHROMA_COLLECTION", COLLECTION_NAME)
    else:
        collection_name = None

    if args.download_drive:
        if not DRIVE_FOLDER_ID:
            logging.error("Falta DRIVE_FOLDER_ID en .env")
            exit(1)
        logging.info("Descargando actas desde Google Drive...")
        success = download_folder_from_drive(DRIVE_FOLDER_ID, RAW_DOCS_PATH)
        if not success:
            exit(1)
    elif args.download_raw:
        if not BUCKET_NAME:
            logging.error("Falta GCP_BUCKET_NAME en .env")
            exit(1)
        logging.info("Descargando actas crudas desde GCP Bucket...")
        download_folder_from_gcs(BUCKET_NAME, "raw_docs", RAW_DOCS_PATH)

    logging.info(
        "Iniciando procesamiento local (chunk_strategy=%s)...",
        args.chunk_strategy,
    )
    os.makedirs(RAW_DOCS_PATH, exist_ok=True)

    if args.chunk_strategy == "both":
        run_indexing(RAW_DOCS_PATH, CHROMA_PATH, "both")
    else:
        documents = process_all_documents(RAW_DOCS_PATH, chunk_strategy=args.chunk_strategy)
        if documents:
            index_to_chroma(documents, CHROMA_PATH, collection_name)
        else:
            logging.warning("El pipeline se detuvo por falta de documentos.")
            exit(0)

    if args.upload_db:
        if not BUCKET_NAME:
            logging.error("Falta GCP_BUCKET_NAME en .env")
            exit(1)
        logging.info("Subiendo ChromaDB a GCP...")
        upload_folder_to_gcs(BUCKET_NAME, CHROMA_PATH, "chroma_db")
        logging.info("✅ Upload completado con éxito.")

if __name__ == "__main__":
    main()
